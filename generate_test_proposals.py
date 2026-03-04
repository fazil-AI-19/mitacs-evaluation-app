"""
Generate 4 test Mitacs Accelerate research proposals as .docx files.

Scenarios:
  1. Strong proposal (expected: Accept) — AI-driven predictive maintenance
  2. Moderate proposal (expected: Revise & Resubmit) — Carbon footprint analytics
  3. Weak proposal (expected: Reject) — Vague software upgrade project
  4. Indigenous-involvement proposal (expected: Revise & Resubmit) — Indigenous health data platform
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_DIR = Path(__file__).parent / "test_proposals"
OUTPUT_DIR.mkdir(exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")


def add_section_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True


# ─────────────────────────────────────────────────────────────────────────────
# Proposal 1 — Strong (Accept)
# AI-Driven Predictive Maintenance for Automotive Manufacturing
# ─────────────────────────────────────────────────────────────────────────────

def proposal_1(doc: Document) -> None:
    doc.add_heading("Mitacs Accelerate Research Proposal", 1)
    doc.add_paragraph(
        "Applicant: Dr. Sarah Chen | Institution: University of Waterloo | "
        "Partner: AutoPrecision Manufacturing Ltd."
    )

    add_heading(doc, "1.1 Project Summary")
    add_body(doc,
        "AutoPrecision Manufacturing Ltd. operates three stamping plants across southern Ontario "
        "producing precision automotive components. Unplanned equipment downtime currently costs "
        "the company approximately $2.4 million per year in lost throughput and emergency repairs. "
        "Existing maintenance schedules are time-based and do not adapt to actual machine condition.\n\n"
        "This project will develop a machine-learning-based predictive maintenance system that "
        "ingests vibration, temperature, current draw, and acoustic signals from 42 CNC milling "
        "centres and hydraulic presses to predict failure events 72–96 hours in advance with "
        "greater than 90% precision. The intern will train anomaly detection and time-series "
        "classification models on 18 months of historical sensor logs, validate predictions "
        "against ground-truth maintenance records, and deploy the system as an edge-inference "
        "module integrated with AutoPrecision's existing SCADA network.\n\n"
        "Expected benefits: 35–45% reduction in unplanned downtime, $800K–$1.2M annual savings "
        "for the partner, one new hire in data engineering, and a generalizable methodology "
        "publishable in a peer-reviewed journal on intelligent manufacturing."
    )

    add_heading(doc, "1.2 Project Context")
    add_section_label(doc, "Research Problem:")
    add_body(doc,
        "Can multivariate sensor fusion and deep time-series models predict catastrophic failure "
        "in CNC machining centres and hydraulic presses 72–96 hours before occurrence with "
        "precision ≥ 90% and recall ≥ 85%, reducing unplanned downtime by at least 35%?"
    )
    add_section_label(doc, "Background and Review of Relevant Prior Work:")
    add_body(doc,
        "Predictive maintenance (PdM) has emerged as a critical capability in Industry 4.0 "
        "manufacturing. Traditional time-based and condition-based maintenance strategies suffer "
        "from either over-maintenance (unnecessary part replacement) or under-maintenance "
        "(missed failure events). Data-driven PdM addresses this by learning degradation patterns "
        "directly from sensor streams [1].\n\n"
        "Vibration analysis using Fast Fourier Transform (FFT) and wavelet decomposition has been "
        "the dominant technique for rotating machinery fault detection since the 1990s [2]. "
        "However, FFT-based approaches require domain expertise to select meaningful frequency "
        "bands and fail to model temporal dependencies between fault signatures across multiple "
        "machines. Recent work by Zhao et al. [3] demonstrated that Long Short-Term Memory (LSTM) "
        "networks outperform traditional statistical process control charts on bearing degradation "
        "datasets (PRONOSTIA, CWRU) by 18–22% in F1 score.\n\n"
        "Transformer architectures, originally developed for natural language processing, have "
        "shown strong results on multivariate time-series anomaly detection. Wu et al. [4] "
        "introduced Anomaly Transformer, achieving state-of-the-art on the MSL and SMAP satellite "
        "telemetry benchmarks. Applying these architectures to manufacturing sensor fusion remains "
        "an open research problem: industrial datasets differ in sampling rate heterogeneity, "
        "significant class imbalance (failure events are rare), and noisy labeling of ground-truth "
        "failure times [5].\n\n"
        "AutoPrecision's dataset (18 months, 42 machines, 6 sensor modalities, ~2.1 billion "
        "readings) represents a unique opportunity to advance the field. No published work has "
        "benchmarked transformer-based anomaly detection under the specific constraints of "
        "automotive press-line environments (coolant contamination, variable spindle loads, "
        "coordinated multi-machine workpieces). This project fills that gap and contributes a "
        "new benchmark dataset and evaluation protocol to the research community.\n\n"
        "References:\n"
        "[1] Lee, J. et al. (2014). Service innovation and smart analytics for Industry 4.0 and "
        "big data environment. Procedia CIRP, 16, 3–8.\n"
        "[2] Randall, R.B. (2011). Vibration-based Condition Monitoring. Wiley.\n"
        "[3] Zhao, R. et al. (2019). Machine health monitoring using local feature-based gated "
        "recurrent unit networks. IEEE Transactions on Industrial Electronics, 65(2).\n"
        "[4] Wu, H. et al. (2021). Anomaly Transformer: Time Series Anomaly Detection with "
        "Association Discrepancy. ICLR 2022.\n"
        "[5] Ran, Y. et al. (2019). A survey of predictive maintenance: Systems, purposes "
        "and approaches. arXiv:1912.07383."
    )

    add_heading(doc, "1.3 Overall Project Objectives")
    add_body(doc,
        "Develop, validate, and deploy a transformer-based multivariate predictive maintenance "
        "system for AutoPrecision's CNC and hydraulic press fleet that predicts failure events "
        "72–96 hours in advance with precision ≥ 90%, recall ≥ 85%, and reduces unplanned "
        "downtime by ≥ 35% as measured over a six-month deployment window."
    )

    add_heading(doc, "1.4 Project Sub-Objectives & Timeline")
    add_body(doc,
        "Duration: 12 months | Intern: MSc 1 (Machine Learning), MSc 2 (Systems Integration)\n\n"
        "Sub-Objective 1 — Data pipeline and preprocessing (Months 1–3, MSc 1 & MSc 2):\n"
        "  Task 1.1: Audit raw sensor data; identify gaps, calibration drift, and labeling "
        "inconsistencies (Months 1–2)\n"
        "  Task 1.2: Implement streaming ETL pipeline with Kafka and Apache Spark; normalize "
        "multi-rate signals to 1 Hz unified window (Months 2–3)\n\n"
        "Sub-Objective 2 — Model development and benchmarking (Months 3–7, MSc 1):\n"
        "  Task 2.1: Implement baseline models (isolation forest, LSTM autoencoder) (Month 3–4)\n"
        "  Task 2.2: Implement and tune Anomaly Transformer on AutoPrecision dataset (Months 4–6)\n"
        "  Task 2.3: Conduct ablation studies; compare single-modal vs. fused sensor inputs (Month 6–7)\n\n"
        "Sub-Objective 3 — Edge deployment and SCADA integration (Months 7–10, MSc 2):\n"
        "  Task 3.1: Quantize and optimize model for NVIDIA Jetson NX edge deployment (Months 7–8)\n"
        "  Task 3.2: Develop REST API and OPC-UA adapter for SCADA integration (Months 8–9)\n"
        "  Task 3.3: Pilot deployment on 8 machines; validate real-time inference latency < 500ms (Months 9–10)\n\n"
        "Sub-Objective 4 — Full deployment and evaluation (Months 10–12, MSc 1 & MSc 2):\n"
        "  Task 4.1: Roll out to all 42 machines; monitor downtime KPIs (Months 10–11)\n"
        "  Task 4.2: Prepare journal manuscript and internal technical report (Months 11–12)"
    )

    add_heading(doc, "1.5 Methodology/Approach")
    add_body(doc,
        "The project uses a four-phase methodology: (1) data preparation, (2) model development, "
        "(3) edge deployment, and (4) production validation.\n\n"
        "Data preparation: Raw vibration (accelerometers, 10 kHz), temperature (thermocouples, "
        "1 Hz), current draw (Hall-effect sensors, 200 Hz), acoustic emission (MEMS microphones, "
        "20 kHz), and spindle load (encoder, 100 Hz) will be synchronized to a unified 1-second "
        "window using Apache Spark structured streaming on AutoPrecision's on-premises Hadoop "
        "cluster. Ground-truth failure labels will be extracted from the SAP PM work-order system "
        "and cross-referenced with operator logbooks. Class imbalance (estimated 1:800 "
        "failure:normal) will be addressed using a combination of synthetic oversampling (SMOTE "
        "on windowed features) and cost-sensitive loss weighting.\n\n"
        "Model development: Anomaly Transformer will be adapted for heterogeneous sensor fusion "
        "by introducing a modality-specific embedding layer followed by cross-modal attention. "
        "The model will be trained on 70% of the timeline (earliest 13 months), validated on "
        "the next 3 months (temporal holdout), and tested on the final 2 months. Hyperparameter "
        "search will use Bayesian optimization (Optuna) with 5-fold cross-validation over the "
        "training window to prevent lookahead bias.\n\n"
        "Risk mitigation: If transformer training is computationally prohibitive on available "
        "GPU resources (2× NVIDIA A100 provided by University of Waterloo's WatCloud), we will "
        "fall back to a patch-based LSTM variant that has lower memory footprint with comparable "
        "accuracy on comparable industrial datasets. Model drift will be monitored using "
        "population stability index (PSI) on incoming sensor distributions, with automated "
        "retraining triggered when PSI > 0.2.\n\n"
        "Validation: Success criteria are defined prospectively — precision ≥ 90%, recall ≥ 85% "
        "on the held-out test period — and confirmed by downtime KPI improvement over a 6-month "
        "deployment window documented by AutoPrecision's maintenance manager."
    )

    add_heading(doc, "1.6 Deliverables")
    add_body(doc,
        "✓ Peer-reviewed journal article (IEEE Transactions on Industrial Informatics, target)\n"
        "✓ Conference presentation (IISE Annual Conference 2025 or PHM Society)\n"
        "✓ Production-grade predictive maintenance software deployed on AutoPrecision SCADA\n"
        "✓ Technical report documenting dataset, benchmark results, and deployment guide\n"
        "✓ New research tools: open-source benchmark dataset and evaluation harness (GitHub)\n"
        "✓ One new data engineering hire at AutoPrecision\n"
        "✓ Two MSc theses (one ML, one systems integration)"
    )

    add_heading(doc, "1.7 Partner Interaction")
    add_body(doc,
        "Interaction mode: Hybrid — on-site three days per week at AutoPrecision's Brantford "
        "plant during data pipeline and deployment phases; remote during modelling phases.\n\n"
        "Activities: MSc 1 will work with AutoPrecision's data engineering team to access and "
        "validate historical sensor archives. MSc 2 will work with plant maintenance engineers "
        "to integrate the edge module with SCADA. Both interns will present monthly progress "
        "updates to the VP of Operations and maintenance supervisor.\n\n"
        "Resources provided: Full access to SAP PM system, SCADA historian, on-site server "
        "rack for edge device testing, and a dedicated workstation. The maintenance manager "
        "(P.Eng.) will serve as the industrial co-supervisor.\n\n"
        "Academic supervision: Dr. Sarah Chen (supervisor) will hold weekly one-hour meetings "
        "with each intern and monthly joint group meetings. Interns will present their work at "
        "the University of Waterloo Systems Design seminar series.\n\n"
        "Career development: Interns gain hands-on experience in industrial ML deployment, "
        "OPC-UA industrial protocols, cross-functional project management, and collaboration "
        "with P.Eng. maintenance professionals — directly applicable to roles in smart "
        "manufacturing, industrial AI, and digital twin development."
    )

    add_heading(doc, "1.8 Benefit to the Partner Organization and to Canada")
    add_body(doc,
        "✓ R&D to improve a product, process or market innovation\n"
        "✓ Improved productivity for partner organization\n"
        "✓ Contribution to the development of skilled talent in Canada\n\n"
        "AutoPrecision will reduce unplanned downtime by an estimated 35–45%, translating to "
        "$800K–$1.2M annual savings. The system is architected for replication across "
        "AutoPrecision's three plants and is designed to be white-labelled for sale to Tier 1 "
        "automotive suppliers — a potential new revenue line. Two MSc graduates with rare "
        "industrial ML deployment skills will enter Canada's advanced manufacturing workforce.\n\n"
        "IP arrangement: AutoPrecision holds exclusive commercialization rights to the deployed "
        "system. The underlying algorithms and benchmark dataset remain jointly owned and will "
        "be published openly per Mitacs's open science guidelines. No export-controlled "
        "technologies or foreign entity partnerships are involved."
    )

    add_heading(doc, "1.9 Indigenous Community Involvement or Impact")
    add_body(doc,
        "This project does not involve or impact Indigenous communities, traditional knowledge, "
        "or Indigenous participants. Section 1.9 is not applicable."
    )

    add_heading(doc, "1.10 Relationship with Past/Present Mitacs Projects")
    add_body(doc, "No — neither the academic supervisor nor the partner organization has current "
        "or upcoming submissions to Mitacs. Dr. Chen completed one prior Mitacs Accelerate "
        "project (IT15234) on supply chain optimization (2021–2022); the current project is "
        "independent and addresses a distinct technical domain (prognostics vs. logistics)."
    )

    add_heading(doc, "1.11 References")
    add_body(doc,
        "[1] Lee, J. et al. (2014). Procedia CIRP, 16, 3–8.\n"
        "[2] Randall, R.B. (2011). Vibration-based Condition Monitoring. Wiley.\n"
        "[3] Zhao, R. et al. (2019). IEEE Transactions on Industrial Electronics, 65(2).\n"
        "[4] Wu, H. et al. (2021). Anomaly Transformer. ICLR 2022.\n"
        "[5] Ran, Y. et al. (2019). arXiv:1912.07383."
    )


# ─────────────────────────────────────────────────────────────────────────────
# Proposal 2 — Moderate (Revise & Resubmit)
# Carbon Footprint Analytics Platform for Retail Supply Chains
# ─────────────────────────────────────────────────────────────────────────────

def proposal_2(doc: Document) -> None:
    doc.add_heading("Mitacs Accelerate Research Proposal", 1)
    doc.add_paragraph(
        "Applicant: Dr. Marco Bianchi | Institution: Dalhousie University | "
        "Partner: GreenRetail Solutions Inc."
    )

    add_heading(doc, "1.1 Project Summary")
    add_body(doc,
        "GreenRetail Solutions Inc. is a Halifax-based sustainability consulting firm that helps "
        "mid-market retailers measure and reduce their Scope 1, 2, and 3 greenhouse gas (GHG) "
        "emissions. Current tools require significant manual data entry from suppliers and "
        "produce static annual reports that are too slow and costly for large retail networks.\n\n"
        "This project will develop a semi-automated carbon analytics platform that integrates "
        "purchase-order transaction data, shipping manifests, and supplier emission factors to "
        "produce near-real-time Scope 3 estimates at the product-category level. The research "
        "challenge is developing a robust imputation framework for missing supplier data, which "
        "affects roughly 40% of a typical retail supply chain.\n\n"
        "Expected outcomes include a functional prototype validated against two GreenRetail "
        "clients, a peer-reviewed publication on supply chain emission imputation methods, "
        "and a scalable platform that GreenRetail plans to commercialize as a SaaS product."
    )

    add_heading(doc, "1.2 Project Context")
    add_section_label(doc, "Research Problem:")
    add_body(doc,
        "How can missing supplier emission factor data (affecting ~40% of retail procurement "
        "records) be reliably imputed using publicly available databases and machine learning, "
        "to produce Scope 3 Category 1 estimates within ±15% of actual values?"
    )
    add_section_label(doc, "Background and Review of Relevant Prior Work:")
    add_body(doc,
        "Corporate GHG accounting is governed by the GHG Protocol Corporate Standard [1] and "
        "its Scope 3 supplement [2]. Scope 3 Category 1 (purchased goods and services) is "
        "typically the largest emission source for retailers, often representing 70–80% of "
        "total lifecycle emissions [3]. Yet it remains the hardest to measure accurately because "
        "it requires emission factor data from thousands of suppliers.\n\n"
        "Two primary approaches exist in the literature. The spend-based method uses economic "
        "input-output (EIO) models (e.g., USEEIO, Exiobase) to estimate emissions per dollar "
        "of expenditure by industry sector [4]. This is tractable at scale but introduces large "
        "uncertainties (±50% or more) because sector-average factors mask firm-level variation. "
        "The supplier-specific method collects actual emission data from each supplier and is far "
        "more accurate but is resource-intensive and yields 40–60% non-response rates in "
        "practice [5].\n\n"
        "Recent work on data fusion for supply chain emission estimation has combined EIO data "
        "with partial supplier disclosures using multiple imputation by chained equations (MICE) "
        "[6] and gradient boosting regressors [7]. These methods reduce uncertainty to ±20–25% "
        "in controlled studies but have not been validated on live retail procurement systems. "
        "The project will extend this line of work using GreenRetail's two client datasets as "
        "held-out validation environments.\n\n"
        "References:\n"
        "[1] GHG Protocol Corporate Standard (2004, revised 2015). WBCSD/WRI.\n"
        "[2] GHG Protocol Scope 3 Standard (2011). WBCSD/WRI.\n"
        "[3] CDP Supply Chain Report (2022). CDP Worldwide.\n"
        "[4] Hertwich, E. & Peters, G. (2009). Carbon footprint of nations. Env. Sci. Tech., 43.\n"
        "[5] Foerstl, K. et al. (2017). Supply chain decarbonization. J. Cleaner Production, 147.\n"
        "[6] van Buuren, S. & Groothuis-Oudshoorn, K. (2011). MICE. J. Statistical Software, 45.\n"
        "[7] Chen, T. & Guestrin, C. (2016). XGBoost. ACM KDD."
    )

    add_heading(doc, "1.3 Overall Project Objectives")
    add_body(doc,
        "Design and validate a machine-learning-based missing-data imputation framework for "
        "Scope 3 Category 1 emissions estimation that achieves ±15% accuracy versus "
        "supplier-disclosed values, and package the framework as a deployable analytics "
        "module integrated with GreenRetail's client reporting pipeline."
    )

    add_heading(doc, "1.4 Project Sub-Objectives & Timeline")
    add_body(doc,
        "Duration: 12 months | Intern: PhD 1 (Environmental Data Science)\n\n"
        "Sub-Objective 1 — Data integration and baseline (Months 1–4):\n"
        "  Task 1.1: Map GreenRetail client transaction schemas to GHG Protocol categories\n"
        "  Task 1.2: Integrate Exiobase v3.8 and EPA USEEIO 2.0 as baseline imputation sources\n"
        "  Task 1.3: Quantify missingness patterns and establish spend-based baseline accuracy\n\n"
        "Sub-Objective 2 — Imputation model development (Months 3–8):\n"
        "  Task 2.1: Implement MICE imputation using publicly available emission databases\n"
        "  Task 2.2: Develop XGBoost ensemble using supplier characteristics as covariates\n"
        "  Task 2.3: Evaluate models on Client A dataset (holdout last 3 months)\n\n"
        "Sub-Objective 3 — Platform development and validation (Months 7–11):\n"
        "  Task 3.1: Implement FastAPI backend with imputation pipeline and report generation\n"
        "  Task 3.2: Validate on Client B dataset\n"
        "  Task 3.3: GreenRetail pilot deployment and user acceptance testing\n\n"
        "Sub-Objective 4 — Dissemination (Months 10–12):\n"
        "  Task 4.1: Submit manuscript to Journal of Cleaner Production\n"
        "  Task 4.2: Present at IEEE SSCI or SETAC North America"
    )

    add_heading(doc, "1.5 Methodology/Approach")
    add_body(doc,
        "The study will use a quantitative research design. Procurement data from two GreenRetail "
        "clients will be anonymized and provided under data sharing agreements. Supplier records "
        "will be linked to Exiobase sectoral emission intensities using NAICS codes.\n\n"
        "We will train imputation models on Client A data (3 years, ~85,000 SKU-supplier pairs) "
        "and evaluate on Client B (2 years, ~60,000 pairs) to assess generalizability. "
        "The primary metric is mean absolute percentage error (MAPE) vs. supplier-disclosed "
        "CDP emission factors for the subset of suppliers with known disclosures.\n\n"
        "One limitation is that only about 18% of suppliers in the combined dataset have "
        "disclosed emission factors, which constrains the ground-truth evaluation. We will "
        "supplement with a sensitivity analysis using synthetic data generated by varying "
        "known supplier characteristics within literature-reported ranges.\n\n"
        "The timeline is somewhat aggressive for a single PhD intern. If the platform development "
        "phase extends, we may defer user acceptance testing to a post-project commercialization "
        "phase that GreenRetail will fund independently."
    )

    add_heading(doc, "1.6 Deliverables")
    add_body(doc,
        "✓ Journal manuscript (Journal of Cleaner Production)\n"
        "✓ Conference presentation\n"
        "✓ Functional analytics platform prototype\n"
        "✓ Technical report for GreenRetail clients\n"
        "✓ PhD thesis chapter"
    )

    add_heading(doc, "1.7 Partner Interaction")
    add_body(doc,
        "Interaction mode: Hybrid — on-site at GreenRetail's Halifax office one day per week; "
        "remote for data analysis and writing phases.\n\n"
        "The intern will work directly with GreenRetail's analytics lead and two sustainability "
        "consultants. Monthly steering committee meetings will include Dr. Bianchi, the intern, "
        "and GreenRetail's CEO.\n\n"
        "Academic supervision: Weekly supervision meetings with Dr. Bianchi; the intern will "
        "also participate in Dalhousie's Environmental Informatics lab group meeting bi-weekly.\n\n"
        "Career development: The intern will gain experience in applied sustainability analytics, "
        "client-facing project management, and Python-based data product development."
    )

    add_heading(doc, "1.8 Benefit to the Partner Organization and to Canada")
    add_body(doc,
        "✓ R&D to improve a product, process or market innovation\n"
        "✓ Contribution to the development of skilled talent in Canada\n\n"
        "GreenRetail will commercialize the imputation framework as a SaaS module, reducing "
        "client onboarding from 3 months to 2 weeks. This addresses a $4B Canadian market "
        "for corporate sustainability reporting services. The PhD intern will develop skills "
        "directly applicable to environmental data science roles in the growing Canadian "
        "clean-tech sector."
    )

    add_heading(doc, "1.9 Indigenous Community Involvement or Impact")
    add_body(doc,
        "This project does not involve or impact Indigenous communities. Section 1.9 is "
        "not applicable."
    )

    add_heading(doc, "1.10 Relationship with Past/Present Mitacs Projects")
    add_body(doc,
        "No current or upcoming concurrent submissions. Dr. Bianchi had a prior Mitacs "
        "Accelerate project (IT14897) on life cycle assessment methodology (2020–2021). "
        "The current project builds on LCA methodology from that work but targets a new "
        "application domain (retail supply chains vs. construction materials) and "
        "introduces a new research contribution (ML-based imputation)."
    )

    add_heading(doc, "1.11 References")
    add_body(doc,
        "[1] GHG Protocol Corporate Standard (2004, revised 2015).\n"
        "[2] GHG Protocol Scope 3 Standard (2011).\n"
        "[3] CDP Supply Chain Report (2022).\n"
        "[4] Hertwich & Peters (2009). Env. Sci. Tech., 43.\n"
        "[5] Foerstl et al. (2017). J. Cleaner Production, 147.\n"
        "[6] van Buuren & Groothuis-Oudshoorn (2011). J. Statistical Software, 45.\n"
        "[7] Chen & Guestrin (2016). XGBoost. ACM KDD."
    )


# ─────────────────────────────────────────────────────────────────────────────
# Proposal 3 — Weak (Reject)
# Upgrading Legacy Database Software for a Financial Services Firm
# ─────────────────────────────────────────────────────────────────────────────

def proposal_3(doc: Document) -> None:
    doc.add_heading("Mitacs Accelerate Research Proposal", 1)
    doc.add_paragraph(
        "Applicant: Dr. Kevin Park | Institution: Lakehead University | "
        "Partner: Northern Capital Financial Services Ltd."
    )

    add_heading(doc, "1.1 Project Summary")
    add_body(doc,
        "Northern Capital Financial Services Ltd. is a Thunder Bay-based financial planning "
        "firm that uses a legacy Microsoft Access database to manage client portfolios. The "
        "system is slow and difficult to maintain. This project will help Northern Capital "
        "upgrade their database to a modern SQL system and improve their reporting capabilities.\n\n"
        "The intern will assist Northern Capital's IT team in migrating data from Access to "
        "PostgreSQL and building some new report templates. We expect this will make the company "
        "more efficient and save staff time on manual reporting tasks."
    )

    add_heading(doc, "1.2 Project Context")
    add_section_label(doc, "Research Problem:")
    add_body(doc,
        "How can Northern Capital migrate their legacy database to PostgreSQL in a way that "
        "minimizes disruption to daily operations?"
    )
    add_section_label(doc, "Background and Review of Relevant Prior Work:")
    add_body(doc,
        "Database migration is a well-known challenge in IT. Many companies face issues when "
        "moving from old systems to new ones. Microsoft Access is widely used by small businesses "
        "but has scalability limitations [1]. PostgreSQL is an open-source relational database "
        "that offers better performance [2].\n\n"
        "Some companies have had success migrating to cloud databases like AWS RDS. Others have "
        "used ETL tools to automate data transfer. There is general consensus that having a "
        "clear migration plan reduces risk [3].\n\n"
        "References:\n"
        "[1] Microsoft (2023). Access product documentation.\n"
        "[2] PostgreSQL Global Development Group (2023). PostgreSQL 15 documentation.\n"
        "[3] Various industry best practices."
    )

    add_heading(doc, "1.3 Overall Project Objectives")
    add_body(doc,
        "Migrate Northern Capital's client database from Microsoft Access to PostgreSQL and "
        "build new reporting dashboards."
    )

    add_heading(doc, "1.4 Project Sub-Objectives & Timeline")
    add_body(doc,
        "Duration: 6 months | Intern: Undergraduate 1 (Computer Science)\n\n"
        "Month 1–2: Understand the existing database and document its structure.\n"
        "Month 3–4: Migrate the data to PostgreSQL.\n"
        "Month 5: Build report templates.\n"
        "Month 6: Testing and handover to Northern Capital IT staff."
    )

    add_heading(doc, "1.5 Methodology/Approach")
    add_body(doc,
        "The intern will first meet with Northern Capital staff to understand their needs. "
        "Then they will create a migration plan. After that, the data will be transferred "
        "to the new system. The intern will use standard database migration tools and "
        "SQL scripts. Report templates will be built in Microsoft Power BI.\n\n"
        "Testing will be done to make sure the new system works correctly before handover."
    )

    add_heading(doc, "1.6 Deliverables")
    add_body(doc,
        "✓ Migrated PostgreSQL database\n"
        "✓ Power BI report templates\n"
        "✓ Technical documentation for Northern Capital IT staff"
    )

    add_heading(doc, "1.7 Partner Interaction")
    add_body(doc,
        "The intern will work on-site at Northern Capital's Thunder Bay office. They will "
        "interact with the IT manager and two administrative staff. Dr. Park will check in "
        "with the intern monthly by phone."
    )

    add_heading(doc, "1.8 Benefit to the Partner Organization and to Canada")
    add_body(doc,
        "✓ Improved productivity for partner organization\n\n"
        "Northern Capital will save time on manual reporting. Staff will be able to access "
        "client information more quickly. The undergraduate intern will gain work experience."
    )

    add_heading(doc, "1.9 Indigenous Community Involvement or Impact")
    add_body(doc, "Not applicable.")

    add_heading(doc, "1.10 Relationship with Past/Present Mitacs Projects")
    add_body(doc, "No previous Mitacs projects.")

    add_heading(doc, "1.11 References")
    add_body(doc,
        "[1] Microsoft (2023). Access product documentation.\n"
        "[2] PostgreSQL Global Development Group (2023).\n"
        "[3] Various industry best practices."
    )


# ─────────────────────────────────────────────────────────────────────────────
# Proposal 4 — Indigenous Involvement (Revise & Resubmit)
# Culturally Adapted Digital Health Navigation Platform for First Nations Communities
# ─────────────────────────────────────────────────────────────────────────────

def proposal_4(doc: Document) -> None:
    doc.add_heading("Mitacs Accelerate Research Proposal", 1)
    doc.add_paragraph(
        "Applicant: Dr. Anika Patel | Institution: University of Manitoba | "
        "Partner: Turtle Island Wellness Collective"
    )

    add_heading(doc, "1.1 Project Summary")
    add_body(doc,
        "Turtle Island Wellness Collective (TIWC) is an Indigenous-led non-profit serving "
        "six Anishinaabe communities in northern Manitoba. TIWC operates a community health "
        "navigation program that connects community members with provincial health services "
        "but currently relies entirely on phone-based coordination between health navigators "
        "and clients. The existing system creates gaps in follow-up care and burdens navigators "
        "with manual scheduling.\n\n"
        "This project will co-design and develop a culturally adapted digital health navigation "
        "platform that supports secure messaging, appointment booking, and care plan tracking "
        "between navigators and clients, incorporating Indigenous concepts of relational care "
        "and community wellness. The research challenge is to develop a human-computer "
        "interaction (HCI) design methodology that authentically embeds Indigenous values "
        "(particularly the Medicine Wheel framework) into a digital interface, rather than "
        "retrofitting a Western health tech paradigm.\n\n"
        "Expected outcomes: A production-ready mobile application deployed across all six TIWC "
        "communities, a community data governance framework developed in partnership with "
        "community Elders, and a peer-reviewed publication on Indigenous-centered HCI design."
    )

    add_heading(doc, "1.2 Project Context")
    add_section_label(doc, "Research Problem:")
    add_body(doc,
        "How can Indigenous cultural frameworks — specifically the Medicine Wheel's four "
        "dimensions of holistic health — be operationalized as concrete UX design principles "
        "for a digital health navigation tool that is trusted and adopted by Anishinaabe "
        "community members in northern Manitoba?"
    )
    add_section_label(doc, "Background and Review of Relevant Prior Work:")
    add_body(doc,
        "Indigenous peoples in Canada face significant health disparities, with First Nations "
        "communities experiencing higher rates of chronic disease, lower life expectancy, and "
        "limited access to specialist care [1]. A key barrier is discontinuity of care: "
        "community members are often referred to distant urban centres without adequate "
        "follow-up infrastructure [2].\n\n"
        "Health navigation programs — in which trained community members serve as liaisons "
        "between patients and the health system — have demonstrated significant effectiveness "
        "in improving access and outcomes for Indigenous populations [3]. Digital augmentation "
        "of navigation programs has the potential to extend navigator capacity, but existing "
        "health apps are designed for Western individualistic health models and employ "
        "surveillance-oriented architectures that conflict with Indigenous data sovereignty "
        "principles [4].\n\n"
        "The OCAP® (Ownership, Control, Access, and Possession) principles, established by the "
        "First Nations Information Governance Centre, provide a framework for data governance "
        "in Indigenous research contexts [5]. Recent scholarship in Indigenous HCI (e.g., "
        "Eglash et al. on community-based design [6]; Bidwell et al. on rural Indigenous "
        "technology adoption [7]) emphasizes co-design processes led by community knowledge "
        "holders as essential for technology acceptance and sustained use.\n\n"
        "The Medicine Wheel, a framework central to many Plains and Great Lakes Algonquian "
        "traditions, conceptualizes wellness holistically across physical, mental, emotional, "
        "and spiritual dimensions [8]. Operationalizing this as UX design guidance — for "
        "navigation flows, notification logic, and interface language — represents a novel "
        "contribution to both Indigenous HCI and digital health informatics.\n\n"
        "References:\n"
        "[1] FNIGC (2018). National Report of the First Nations Regional Health Survey. Ottawa.\n"
        "[2] Lavoie, J. et al. (2010). Have investments in on-reserve health services "
        "and initiatives promoting community control paid off? Social Science & Medicine, 70.\n"
        "[3] Panagiotopoulos, I. et al. (2014). Patient navigation programs. Systematic review, "
        "J. General Internal Medicine, 29.\n"
        "[4] Kukutai, T. & Taylor, J. (Eds.) (2016). Indigenous Data Sovereignty. ANU Press.\n"
        "[5] FNIGC (2014). Ownership, Control, Access and Possession (OCAP®). Ottawa.\n"
        "[6] Eglash, R. (1999). African Fractals. Rutgers University Press.\n"
        "[7] Bidwell, N. (2016). Moving the centre to design social media in rural Africa. "
        "AI & Society, 31.\n"
        "[8] Absolon, K. (2011). Kaandossiwin: How We Come to Know. Fernwood Publishing."
    )

    add_heading(doc, "1.3 Overall Project Objectives")
    add_body(doc,
        "Co-design, develop, and deploy a culturally adapted digital health navigation "
        "platform for Turtle Island Wellness Collective that operationalizes Medicine Wheel "
        "wellness principles as HCI design patterns, complies with OCAP® data sovereignty "
        "standards, and achieves ≥ 70% sustained adoption among active health navigation "
        "clients within 3 months of deployment."
    )

    add_heading(doc, "1.4 Project Sub-Objectives & Timeline")
    add_body(doc,
        "Duration: 18 months | Intern: PhD 1 (Health Informatics)\n\n"
        "Sub-Objective 1 — Community engagement and co-design (Months 1–5, PhD 1):\n"
        "  Task 1.1: Conduct talking circles with community Elders and health navigators to "
        "identify wellness priorities and cultural values for the platform (Months 1–3)\n"
        "  Task 1.2: Develop and validate Medicine Wheel UX design principles with community "
        "feedback (Months 3–5)\n\n"
        "Sub-Objective 2 — Platform development (Months 4–11, PhD 1):\n"
        "  Task 2.1: Develop community data governance framework with TIWC board and Elders "
        "(Months 4–6)\n"
        "  Task 2.2: Build React Native mobile application with end-to-end encrypted "
        "messaging and FHIR-compliant care plan module (Months 5–9)\n"
        "  Task 2.3: Iterative usability testing with 15 navigators and 30 clients using "
        "community-based participatory research methods (Months 8–11)\n\n"
        "Sub-Objective 3 — Deployment and evaluation (Months 11–16, PhD 1):\n"
        "  Task 3.1: Phased rollout to all six communities with on-site training (Months 11–13)\n"
        "  Task 3.2: Measure adoption, navigator efficiency, and care continuity metrics "
        "(Months 13–16)\n\n"
        "Sub-Objective 4 — Dissemination (Months 15–18, PhD 1):\n"
        "  Task 4.1: Submit manuscript to Journal of the American Medical Informatics "
        "Association (JAMIA) or First Nations Perspectives\n"
        "  Task 4.2: Knowledge translation workshop for TIWC board and partner communities"
    )

    add_heading(doc, "1.5 Methodology/Approach")
    add_body(doc,
        "This project uses Community-Based Participatory Research (CBPR) as its overarching "
        "methodology, following the iterative co-design approach described by Israel et al. [9]. "
        "All research activities have been reviewed and approved by TIWC's community research "
        "ethics committee (letter of support attached).\n\n"
        "Phase 1 (Co-design): Talking circles will be conducted in English and Ojibwe "
        "(with Elder-approved translation) in each of the six communities. Circles will be "
        "facilitated by the PhD intern under the mentorship of TIWC's Cultural Director "
        "(Elder Margaret Swifthawk). Audio recordings will remain under TIWC's custody per "
        "OCAP® principles. Analysis will use an Indigenous grounded theory approach [10].\n\n"
        "Phase 2 (Development): The platform will be built as a React Native app with a "
        "FastAPI backend hosted on a TIWC-controlled server (not cloud-hosted, per community "
        "preference for data locality). End-to-end encryption (Signal Protocol) will be "
        "implemented for all messaging. FHIR R4 will be used for care plan interoperability "
        "with the Manitoba Health provincial system. All data remains under TIWC's ownership.\n\n"
        "Phase 3 (Evaluation): Adoption will be measured using app analytics (session "
        "frequency, feature use). Care continuity will be measured by navigator-reported "
        "follow-up completion rates compared to a 6-month pre-implementation baseline. "
        "A mixed-methods evaluation will include structured interviews with a sample of "
        "navigators (n=10) and clients (n=20) at 3 months post-launch.\n\n"
        "The PhD intern has limited prior experience with Indigenous research methodologies. "
        "To address this, TIWC's Cultural Director will provide structured mentorship "
        "throughout the project, and the intern will complete the First Nations University "
        "of Canada's online certificate in Indigenous Research Methods (Months 1–3). "
        "This training plan has been approved by Elder Swifthawk and the TIWC board.\n\n"
        "References:\n"
        "[9] Israel, B. et al. (2005). Methods in Community-Based Participatory Research "
        "for Health. Jossey-Bass.\n"
        "[10] Kovach, M. (2009). Indigenous Methodologies. University of Toronto Press."
    )

    add_heading(doc, "1.6 Deliverables")
    add_body(doc,
        "✓ Production React Native application deployed across six TIWC communities\n"
        "✓ TIWC Community Data Governance Framework (community-owned document)\n"
        "✓ Peer-reviewed journal article on Indigenous-centered HCI design\n"
        "✓ Knowledge translation workshop and community report (in accessible language)\n"
        "✓ PhD thesis chapter\n"
        "✓ Open-source Medicine Wheel UX design pattern library (with TIWC approval)"
    )

    add_heading(doc, "1.7 Partner Interaction")
    add_body(doc,
        "Interaction mode: Hybrid — the PhD intern will spend 8 weeks on-site across the "
        "six communities (distributed across Phases 1 and 3); remote during development phase. "
        "Travel costs are included in the intern's budget.\n\n"
        "The intern will work directly with TIWC's health navigator team (6 navigators), "
        "Cultural Director (Elder Swifthawk), and CEO (Ms. Leanne Fontaine). Monthly steering "
        "committee meetings will include TIWC board representatives, Dr. Patel, and the intern.\n\n"
        "Academic supervision: Dr. Patel will meet weekly with the intern for the first "
        "6 months and bi-weekly thereafter. The intern will present work at the University "
        "of Manitoba's Global Health seminar and the Annual Indigenous Health Conference.\n\n"
        "Career development: The intern will develop expertise in Indigenous research ethics, "
        "CBPR methodology, mobile health application development, and FHIR health data "
        "standards — a unique interdisciplinary profile for careers in digital health equity."
    )

    add_heading(doc, "1.8 Benefit to the Partner Organization and to Canada")
    add_body(doc,
        "✓ R&D to improve a product, process or market innovation\n"
        "✓ Use of project results to address a societal challenge\n"
        "✓ Contribution to the development of skilled talent in Canada\n"
        "✓ Fostering new collaborations\n\n"
        "TIWC will gain a self-hosted digital platform that extends its navigators' reach, "
        "reduces care gaps, and respects community data sovereignty. The project directly "
        "addresses Indigenous health equity — a priority in the federal government's "
        "reconciliation agenda and Canada's UNDRIP implementation. The open-source UX "
        "pattern library will be available to other Indigenous health organizations.\n\n"
        "IP and data: All data remains owned by TIWC. The application code is jointly owned "
        "with a perpetual license granted to TIWC. No IP transfer to a non-Indigenous "
        "commercial entity. No export-controlled technologies involved."
    )

    add_heading(doc, "1.9 Indigenous Community Involvement or Impact")
    add_body(doc,
        "This project directly involves and serves six Anishinaabe communities in northern "
        "Manitoba through Turtle Island Wellness Collective. The following addresses the four "
        "required elements:\n\n"
        "Community support and respect: TIWC's board of directors voted unanimously to "
        "partner on this project (resolution attached). Elder Margaret Swifthawk has provided "
        "a letter of support and agreed to serve as Cultural Director for the project. "
        "Community input on research priorities was gathered through informal sharing circles "
        "held by TIWC in fall 2023 prior to this submission.\n\n"
        "Collaborative research practices: All research activities will follow CBPR principles. "
        "The PhD intern will be jointly supervised by Dr. Patel (academic) and Elder Swifthawk "
        "(cultural). Community members will participate in all design phases as co-designers, "
        "not merely research subjects. Community co-authorship is planned for the journal "
        "publication.\n\n"
        "Access and use governance: A formal Community Data Governance Framework will be "
        "developed in Sub-Objective 2.1, co-authored by the TIWC board and Elders. All "
        "platform data will be stored on TIWC-owned servers within community jurisdiction. "
        "OCAP® principles govern all data collection, storage, and use. Raw data will not "
        "leave TIWC custody without explicit band council resolution authorization.\n\n"
        "Team experience and expertise: Dr. Patel has 5 years of experience in Indigenous "
        "health research and has completed CIHR's Indigenous Research training program (2021). "
        "The PhD intern has limited experience in Indigenous research methods; this is "
        "addressed through the First Nations University certificate (Months 1–3) and "
        "Elder Swifthawk's direct mentorship throughout the project. Two TIWC health "
        "navigators with lived experience will serve as community research assistants."
    )

    add_heading(doc, "1.10 Relationship with Past/Present Mitacs Projects")
    add_body(doc,
        "Yes — Dr. Patel has a concurrent Mitacs Accelerate project (IT17852) on Indigenous "
        "mental health data infrastructure for a different partner (Sioux Lookout First "
        "Nations Health Authority). The two projects are related in that both address "
        "Indigenous digital health, but are distinct in community context (northern Manitoba "
        "vs. northwestern Ontario), technical focus (navigation/care coordination vs. mental "
        "health screening), and partner organization. Supervision is shared with separate "
        "interns and separate budgets. Dr. Patel has confirmed with the University of Manitoba "
        "research office that this dual supervision is within approved workload limits. "
        "TIWC and SLFNHA have been informed of each other's projects and have no objections."
    )

    add_heading(doc, "1.11 References")
    add_body(doc,
        "[1] FNIGC (2018). National Report of the First Nations Regional Health Survey.\n"
        "[2] Lavoie et al. (2010). Social Science & Medicine, 70.\n"
        "[3] Panagiotopoulos et al. (2014). J. General Internal Medicine, 29.\n"
        "[4] Kukutai & Taylor (2016). Indigenous Data Sovereignty. ANU Press.\n"
        "[5] FNIGC (2014). OCAP® Principles.\n"
        "[6] Eglash (1999). African Fractals. Rutgers.\n"
        "[7] Bidwell (2016). AI & Society, 31.\n"
        "[8] Absolon (2011). Kaandossiwin. Fernwood.\n"
        "[9] Israel et al. (2005). CBPR Methods. Jossey-Bass.\n"
        "[10] Kovach (2009). Indigenous Methodologies. U of Toronto Press."
    )


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

PROPOSALS = [
    ("proposal_1_strong_accept_predictive_maintenance.docx", proposal_1,
     "Strong | Expected: Accept"),
    ("proposal_2_moderate_revise_carbon_analytics.docx", proposal_2,
     "Moderate | Expected: Revise & Resubmit"),
    ("proposal_3_weak_reject_database_migration.docx", proposal_3,
     "Weak | Expected: Reject"),
    ("proposal_4_indigenous_health_navigation.docx", proposal_4,
     "Indigenous involvement | Expected: Revise & Resubmit"),
]

for filename, fill_fn, label in PROPOSALS:
    doc = Document()
    fill_fn(doc)
    out_path = OUTPUT_DIR / filename
    doc.save(out_path)
    print(f"✓ {out_path.name}  [{label}]")

print(f"\nAll proposals saved to: {OUTPUT_DIR}")
