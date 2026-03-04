from typing import Optional, Type

from crewai.tools import BaseTool
from docx import Document
from pydantic import BaseModel, Field


class _DocxReadInput(BaseModel):
    query: Optional[str] = Field(default=None, description="Search query (ignored; full text returned)")


class DocxReadTool(BaseTool):
    """Reads the full text of a .docx file and returns it as plain text."""

    name: str = "Read DOCX Document"
    description: str = (
        "Reads and returns the complete text content of the proposal DOCX file. "
        "Use this to access the full proposal text."
    )
    docx_path: str = ""
    args_schema: Type[BaseModel] = _DocxReadInput

    def _run(self, query: Optional[str] = None) -> str:
        try:
            doc = Document(self.docx_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            return "\n\n".join(paragraphs) if paragraphs else "(empty document)"
        except Exception as e:
            return f"Error reading document: {e}"
